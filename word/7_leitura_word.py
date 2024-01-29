from docx import Document

documento = Document('sampleEnchanced.docx')
# Como obter o índice + texto
# for index, item in enumerate(documento.paragraphs):
#     print(index)
#     print(item.text)

documento.paragraphs[2].text = 'The population of macacos in alaska is rising'

for index, run in enumerate(documento.paragraphs[4].runs):
    print(str(index) + ' ' + run.text)

documento.paragraphs[4].runs[0].text = 'Um novo parágrafo '
documento.paragraphs[4].runs[1].text = 'um texto super importante '
documento.paragraphs[4].runs[2].text = 'outro parágrafo super importante'

documento.save('demo.docx')